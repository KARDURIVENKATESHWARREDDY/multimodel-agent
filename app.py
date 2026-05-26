import os
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import cv2
import docx
import numpy as np
import pdfplumber
import pytesseract
import streamlit as st
from PIL import Image
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_core.prompts import ChatPromptTemplate
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


load_dotenv()


@dataclass
class ExtractedSource:
    source_name: str
    source_type: str
    text: str


def configure_tracing() -> None:
    if os.getenv("LANGSMITH_API_KEY") and os.getenv("LANGSMITH_TRACING", "true").lower() == "true":
        os.environ["LANGCHAIN_TRACING_V2"] = "true"
        os.environ["LANGCHAIN_API_KEY"] = os.getenv("LANGSMITH_API_KEY", "")
        os.environ["LANGCHAIN_PROJECT"] = os.getenv("LANGSMITH_PROJECT", "multimodal-document-analyzer")


def configure_tesseract() -> None:
    tesseract_cmd = os.getenv("TESSERACT_CMD", "").strip()
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd


def get_llm():
    provider = os.getenv("LLM_PROVIDER", "openai").lower().strip()

    if provider == "openai":
        from langchain_openai import ChatOpenAI

        return ChatOpenAI(model=os.getenv("OPENAI_MODEL", "gpt-4o-mini"), temperature=0)

    if provider == "google":
        from langchain_google_genai import ChatGoogleGenerativeAI

        return ChatGoogleGenerativeAI(model=os.getenv("GOOGLE_MODEL", "gemini-1.5-flash"), temperature=0)

    if provider == "groq":
        from langchain_groq import ChatGroq

        return ChatGroq(model=os.getenv("GROQ_MODEL", "llama-3.1-70b-versatile"), temperature=0)

    if provider == "anthropic":
        from langchain_anthropic import ChatAnthropic

        return ChatAnthropic(model=os.getenv("ANTHROPIC_MODEL", "claude-3-5-sonnet-20240620"), temperature=0)

    raise ValueError("Unsupported LLM_PROVIDER. Use one of: openai, google, groq, anthropic")


def chunk_documents(sources: List[ExtractedSource]) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=200)
    docs: List[Document] = []
    for src in sources:
        chunks = splitter.split_text(src.text)
        for idx, chunk in enumerate(chunks):
            docs.append(
                Document(
                    page_content=chunk,
                    metadata={
                        "source": src.source_name,
                        "source_type": src.source_type,
                        "chunk": idx,
                    },
                )
            )
    return docs


def ocr_image(image_bytes: bytes) -> str:
    image = Image.open(BytesIO(image_bytes)).convert("RGB")
    return pytesseract.image_to_string(image, lang=os.getenv("OCR_LANGUAGE", "eng"))


def extract_pdf(pdf_bytes: bytes) -> str:
    all_text: List[str] = []
    with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                all_text.append(text)
            else:
                page_img = page.to_image(resolution=300).original
                all_text.append(pytesseract.image_to_string(page_img, lang=os.getenv("OCR_LANGUAGE", "eng")))
    return "\n\n".join(all_text)


def extract_docx(docx_bytes: bytes) -> str:
    document = docx.Document(BytesIO(docx_bytes))
    lines = [p.text for p in document.paragraphs if p.text.strip()]
    return "\n".join(lines)


def extract_video(video_bytes: bytes) -> str:
    temp_path = Path("_temp_video_upload.mp4")
    temp_path.write_bytes(video_bytes)

    cap = cv2.VideoCapture(str(temp_path))
    fps = cap.get(cv2.CAP_PROP_FPS) or 24
    interval_seconds = float(os.getenv("VIDEO_FRAME_INTERVAL_SECONDS", "2"))
    frame_interval = max(int(fps * interval_seconds), 1)

    texts: List[str] = []
    frame_idx = 0
    success, frame = cap.read()
    while success:
        if frame_idx % frame_interval == 0:
            rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            pil_img = Image.fromarray(rgb)
            frame_text = pytesseract.image_to_string(pil_img, lang=os.getenv("OCR_LANGUAGE", "eng"))
            if frame_text.strip():
                texts.append(f"[Frame {frame_idx}]\n{frame_text}")
        frame_idx += 1
        success, frame = cap.read()

    cap.release()
    if temp_path.exists():
        temp_path.unlink()

    return "\n\n".join(texts)


def extract_from_uploaded_file(uploaded_file) -> ExtractedSource:
    data = uploaded_file.getvalue()
    name = uploaded_file.name
    mime = uploaded_file.type

    if mime.startswith("image/"):
        text = ocr_image(data)
        return ExtractedSource(name, "image", text)
    if mime == "application/pdf":
        text = extract_pdf(data)
        return ExtractedSource(name, "pdf", text)
    if mime in [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ]:
        text = extract_docx(data)
        return ExtractedSource(name, "docx", text)
    if mime.startswith("video/"):
        text = extract_video(data)
        return ExtractedSource(name, "video", text)

    raise ValueError(f"Unsupported file type: {mime}")


def summarize_documents(llm, sources: List[ExtractedSource]) -> str:
    joined = "\n\n".join([f"SOURCE: {s.source_name} ({s.source_type})\n{s.text}" for s in sources])
    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "You summarize extracted OCR/document text faithfully. If details are uncertain, say so clearly.",
            ),
            (
                "human",
                "Summarize the uploaded content in concise bullet points and include a short section by source.\n\n{content}",
            ),
        ]
    )
    chain = prompt | llm
    result = chain.invoke({"content": joined})
    return result.content if hasattr(result, "content") else str(result)


def retrieve_top_k(question: str, docs: List[Document], k: int = 6) -> List[Document]:
    q_terms = set(question.lower().split())
    scored: List[Tuple[int, Document]] = []
    for d in docs:
        content_terms = set(d.page_content.lower().split())
        score = len(q_terms.intersection(content_terms))
        scored.append((score, d))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [d for s, d in scored[:k] if s > 0] or [d for _, d in scored[:k]]


def answer_grounded(llm, question: str, docs: List[Document]) -> Tuple[str, List[Dict[str, str]]]:
    top_docs = retrieve_top_k(question, docs)
    context = "\n\n".join(
        [
            f"[Source: {d.metadata['source']} | Type: {d.metadata['source_type']} | Chunk: {d.metadata['chunk']}]\n{d.page_content}"
            for d in top_docs
        ]
    )

    system_prompt = (
        "Answer ONLY using the provided context. If context lacks answer, say 'Not found in uploaded documents'. "
        "Cite sources at the end as bullet list with source name and chunk."
    )
    messages = [SystemMessage(content=system_prompt), HumanMessage(content=f"Question: {question}\n\nContext:\n{context}")]
    result = llm.invoke(messages)
    answer = result.content if hasattr(result, "content") else str(result)

    citations = [
        {
            "source": str(d.metadata.get("source", "")),
            "chunk": str(d.metadata.get("chunk", "")),
            "type": str(d.metadata.get("source_type", "")),
        }
        for d in top_docs
    ]
    return answer, citations


def resume_analysis(llm, resume_text: str, job_description: str) -> str:
    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "You are an expert recruiter. Provide grounded analysis only from resume and JD.",
            ),
            (
                "human",
                "Analyze resume match against job description. Provide:\n"
                "1) Match score (0-100) with rationale\n"
                "2) Strengths\n"
                "3) Gaps/missing requirements\n"
                "4) Suggestions to improve alignment\n\n"
                "Resume:\n{resume}\n\nJob Description:\n{jd}",
            ),
        ]
    )
    chain = prompt | llm
    result = chain.invoke({"resume": resume_text, "jd": job_description})
    return result.content if hasattr(result, "content") else str(result)


def to_txt_bytes(text: str) -> bytes:
    return text.encode("utf-8")


def to_docx_bytes(text: str) -> bytes:
    d = docx.Document()
    for line in text.splitlines():
        d.add_paragraph(line)
    stream = BytesIO()
    d.save(stream)
    return stream.getvalue()


def to_pdf_bytes(text: str) -> bytes:
    stream = BytesIO()
    c = canvas.Canvas(stream, pagesize=A4)
    width, height = A4
    y = height - 50
    for raw_line in text.splitlines() or [""]:
        line = raw_line
        if not line.strip():
            y -= 16
            continue
        while len(line) > 110:
            c.drawString(40, y, line[:110])
            line = line[110:]
            y -= 16
            if y < 40:
                c.showPage()
                y = height - 50
        c.drawString(40, y, line)
        y -= 16
        if y < 40:
            c.showPage()
            y = height - 50
    c.save()
    return stream.getvalue()


def export_download_buttons(label_prefix: str, content: str, base_name: str) -> None:
    st.download_button(
        f"{label_prefix} TXT",
        data=to_txt_bytes(content),
        file_name=f"{base_name}.txt",
        mime="text/plain",
    )
    st.download_button(
        f"{label_prefix} DOCX",
        data=to_docx_bytes(content),
        file_name=f"{base_name}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    st.download_button(
        f"{label_prefix} PDF",
        data=to_pdf_bytes(content),
        file_name=f"{base_name}.pdf",
        mime="application/pdf",
    )


def run_document_mode(llm) -> None:
    st.subheader("Multimodal Document Analyzer")
    uploaded_files = st.file_uploader(
        "Upload images, PDFs, DOCX, or videos",
        type=["png", "jpg", "jpeg", "bmp", "tiff", "pdf", "docx", "mp4", "mov", "avi", "mkv"],
        accept_multiple_files=True,
    )

    if "sources" not in st.session_state:
        st.session_state.sources = []
    if "docs" not in st.session_state:
        st.session_state.docs = []
    if "summary" not in st.session_state:
        st.session_state.summary = ""
    if "qa_log" not in st.session_state:
        st.session_state.qa_log = []

    if st.button("Extract Text") and uploaded_files:
        extracted: List[ExtractedSource] = []
        with st.spinner("Extracting text from files..."):
            for f in uploaded_files:
                extracted.append(extract_from_uploaded_file(f))
        st.session_state.sources = extracted
        st.session_state.docs = chunk_documents(extracted)
        st.success(f"Extracted text from {len(extracted)} file(s).")

    if st.session_state.sources:
        with st.expander("Preview Extracted Text"):
            for s in st.session_state.sources:
                st.markdown(f"**{s.source_name}** ({s.source_type})")
                st.text_area(
                    f"Text: {s.source_name}",
                    s.text[:4000],
                    height=150,
                    key=f"preview_{s.source_name}",
                )

        if st.button("Summarize Content"):
            with st.spinner("Summarizing..."):
                st.session_state.summary = summarize_documents(llm, st.session_state.sources)

        if st.session_state.summary:
            st.markdown("### Summary")
            st.write(st.session_state.summary)
            export_download_buttons("Download Summary as", st.session_state.summary, "summary")

        st.markdown("### Ask Questions")
        question = st.text_input("Ask about the uploaded content")
        if st.button("Get Answer") and question.strip():
            with st.spinner("Answering from document context..."):
                answer, citations = answer_grounded(llm, question, st.session_state.docs)
            st.session_state.qa_log.append(
                {
                    "question": question,
                    "answer": answer,
                    "citations": citations,
                }
            )

        if st.session_state.qa_log:
            for i, item in enumerate(st.session_state.qa_log, start=1):
                st.markdown(f"**Q{i}:** {item['question']}")
                st.write(item["answer"])
                st.markdown("**Sources:**")
                for c in item["citations"]:
                    st.markdown(f"- {c['source']} ({c['type']}), chunk {c['chunk']}")
                export_text = (
                    f"Question: {item['question']}\n\nAnswer:\n{item['answer']}\n\n"
                    + "\n".join([f"- {c['source']} ({c['type']}), chunk {c['chunk']}" for c in item["citations"]])
                )
                export_download_buttons(f"Download Q{i} as", export_text, f"qa_{i}")


def run_resume_mode(llm) -> None:
    st.subheader("Resume Analysis Mode")
    resume_file = st.file_uploader("Upload Resume (PDF, DOCX, Image)", type=["pdf", "docx", "png", "jpg", "jpeg"])
    job_description = st.text_area("Paste Job Description", height=220)

    resume_text = ""
    if resume_file:
        fake_wrapper = type("UF", (), {"getvalue": resume_file.getvalue, "name": resume_file.name, "type": resume_file.type})
        extracted = extract_from_uploaded_file(fake_wrapper)
        resume_text = extracted.text
        with st.expander("Resume Text Preview"):
            st.text_area("Extracted Resume Text", resume_text[:5000], height=200)

    if st.button("Analyze Match") and resume_text.strip() and job_description.strip():
        with st.spinner("Analyzing resume against job description..."):
            analysis = resume_analysis(llm, resume_text, job_description)
        st.markdown("### Match Analysis")
        st.write(analysis)
        export_download_buttons("Download Analysis as", analysis, "resume_match_analysis")


def main() -> None:
    st.set_page_config(page_title="Multimodal Document Analyzer", layout="wide")
    st.title("Multimodal Document Analyzer")
    st.caption("Upload files, extract text with OCR, summarize, ask grounded questions, and analyze resume fit.")

    configure_tesseract()
    configure_tracing()

    try:
        llm = get_llm()
    except Exception as e:
        st.error(f"LLM configuration error: {e}")
        st.stop()

    mode = st.radio("Mode", ["Document Analyzer", "Resume Analyzer"], horizontal=True)

    if mode == "Document Analyzer":
        run_document_mode(llm)
    else:
        run_resume_mode(llm)


if __name__ == "__main__":
    main()
